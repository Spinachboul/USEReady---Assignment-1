import streamlit as st
import pandas as pd
import os
from PIL import Image
import pytesseract
from docx import Document
from datetime import datetime
import json
import io
from typing import Dict, List, Any, Tuple
import torch
from transformers import pipeline, AutoTokenizer, AutoModel
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from dateutil import parser
import pickle
from difflib import SequenceMatcher
import re

# Configure page
st.set_page_config(
    page_title="Document Metadata Extractor",
    page_icon="",
    layout="wide"
)

# Title and description
st.title(" Document Metadata Extractor with Training")
st.markdown("Train on your data and extract metadata with improved accuracy")

@st.cache_resource
def load_base_models():
    """Load base AI models with caching"""
    models = {}
    
    try:
        # Question Answering model
        models['qa'] = pipeline(
            "question-answering", 
            model="deepset/roberta-base-squad2",
            return_all_scores=False
        )
        
        models['ner'] = pipeline(
            "ner", 
            model="dbmdz/bert-large-cased-finetuned-conll03-english",
            aggregation_strategy="simple"
        )
        
        models['sentence_transformer'] = pipeline(
            "feature-extraction",
            model="sentence-transformers/all-MiniLM-L6-v2"
        )
        
        st.success("Base AI models loaded successfully!")
        
    except Exception as e:
        st.error(f"Error loading models: {e}")
        models = {}
    
    return models

class TrainableMetadataExtractor:
    def __init__(self, base_models):
        self.base_models = base_models
        self.training_data = []
        self.field_patterns = {}
        self.context_vectors = {}
        self.is_trained = False
        
        # Field extraction strategies
        self.extraction_strategies = {
            "Agreement Value": {
                "questions": [
                    "What is the monetary value of this agreement?",
                    "How much money is involved in this contract?",
                    "What is the total payment amount?",
                    "What is the contract value?",
                    "What is the agreement amount?"
                ],
                "context_keywords": ["value", "amount", "payment", "cost", "price", "fee", "dollar", "USD", "$"]
            },
            "Agreement Start Date": {
                "questions": [
                    "When does this agreement become effective?",
                    "What is the start date of this contract?",
                    "From which date does this agreement begin?",
                    "When does the contract commence?",
                    "What is the effective date?"
                ],
                "context_keywords": ["effective", "commence", "start", "begin", "from", "starting"]
            },
            "Agreement End Date": {
                "questions": [
                    "When does this agreement expire?",
                    "What is the end date of this contract?",
                    "Until when is this agreement valid?",
                    "When does the contract terminate?",
                    "What is the expiration date?"
                ],
                "context_keywords": ["expire", "terminate", "end", "until", "expiration", "conclusion"]
            },
            "Renewal Notice (Days)": {
                "questions": [
                    "How many days notice is required for renewal?",
                    "What is the renewal notice period?",
                    "How many days before renewal must notice be given?",
                    "What is the notification period for renewal?",
                    "How much advance notice is needed?"
                ],
                "context_keywords": ["notice", "renewal", "advance", "prior", "notification", "days"]
            },
            "Party One": {
                "questions": [
                    "Who is the first party in this agreement?",
                    "What is the name of the primary party?",
                    "Who is party one in this contract?",
                    "What company is the first signatory?",
                    "Who is the initial contracting party?"
                ],
                "context_keywords": ["first party", "party one", "primary", "initial", "contracting party"]
            },
            "Party Two": {
                "questions": [
                    "Who is the second party in this agreement?",
                    "What is the name of the secondary party?",
                    "Who is party two in this contract?",
                    "What company is the second signatory?",
                    "Who is the other contracting party?"
                ],
                "context_keywords": ["second party", "party two", "secondary", "other party", "counterpart"]
            }
        }
    
    def extract_text_from_image(self, image_path: str) -> str:
        """Extract text from image using OCR"""
        try:
            image = Image.open(image_path)
            custom_config = r'--oem 3 --psm 6'
            text = pytesseract.image_to_string(image, config=custom_config)
            return text
        except Exception as e:
            st.error(f"Error extracting text from image: {e}")
            return ""
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(docx_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
            
            return "\n".join(text_parts)
        except Exception as e:
            st.error(f"Error extracting text from DOCX: {e}")
            return ""
    
    def train_on_data(self, train_dir: str, train_csv_path: str):
        """Train the extractor on training data"""
        try:
            if not os.path.exists(train_csv_path):
                raise FileNotFoundError(f"Training CSV not found: {train_csv_path}")
            
            # Load training labels
            train_df = pd.read_csv(train_csv_path)
            
            st.info("🔄 Training on provided data...")
            progress_bar = st.progress(0)
            
            self.training_data = []
            
            for idx, row in train_df.iterrows():
                filename = row['File Name'].strip()
                possible_extensions = ['.docx', '.png', '.jpg', '.jpeg']
                filepath = None

                for ext in possible_extensions:
                    cand_path = os.path.join(train_dir, f"{filename}{ext}")
                    if os.path.exists(cand_path):
                        filepath = cand_path
                        break
                
                if not filepath:
                    st.warning(f"File not found for {filename}")
                    continue

                # Extract text
                if filepath.lower().endswith('.docx'):
                    text = self.extract_text_from_docx(filepath)
                else:
                    text = self.extract_text_from_image(filepath)
                
                if text.strip():
                    # Store training example
                    training_example = {
                        'filename': filename,
                        'text': text,
                        'labels': {
                            'Agreement Value': str(row.get('Agreement Value', '')),
                            'Agreement Start Date': str(row.get('Agreement Start Date', '')),
                            'Agreement End Date': str(row.get('Agreement End Date', '')),
                            'Renewal Notice (Days)': str(row.get('Renewal Notice (Days)', '')),
                            'Party One': str(row.get('Party One', '')),
                            'Party Two': str(row.get('Party Two', ''))
                        }
                    }
                    self.training_data.append(training_example)
                
                progress_bar.progress((idx + 1) / len(train_df))
            
            # Learn patterns from training data
            self._learn_extraction_patterns()
            self.is_trained = True
            
            st.success(f"✅ Training completed! Processed {len(self.training_data)} documents")
            
        except Exception as e:
            st.error(f"Training failed: {e}")
            self.is_trained = False
    
    def _learn_extraction_patterns(self):
        """Learn patterns from training data to improve extraction"""
        if not self.training_data:
            return
        
        for field in self.extraction_strategies.keys():
            field_contexts = []
            field_values = []
            
            for example in self.training_data:
                text = example['text']
                label = example['labels'][field]
                
                if label and label.lower() not in ['nan', 'none', '']:
                    # Find context around the label in text
                    context = self._find_context_around_value(text, label)
                    if context:
                        field_contexts.append(context)
                        field_values.append(label)
            
            # Store learned patterns for this field
            self.field_patterns[field] = {
                'contexts': field_contexts,
                'values': field_values
            }
    
    def _find_context_around_value(self, text: str, value: str, window_size: int = 50) -> str:
        """Find context around a specific value in text"""
        try:
            # Simple approach to find context
            text_lower = text.lower()
            value_lower = value.lower()
            
            # Try to find the value or similar text
            pos = text_lower.find(value_lower)
            if pos == -1:
                # Try fuzzy matching for partial matches
                words = text.split()
                for i, word in enumerate(words):
                    if self._similarity_score(word, value) > 0.7:
                        start_idx = max(0, i - window_size//10)
                        end_idx = min(len(words), i + window_size//10)
                        return ' '.join(words[start_idx:end_idx])
                return ""
            
            start = max(0, pos - window_size)
            end = min(len(text), pos + len(value) + window_size)
            return text[start:end]
            
        except:
            return ""
    
    def _similarity_score(self, str1: str, str2: str) -> float:
        """Calculate similarity between two strings"""
        return SequenceMatcher(None, str1.lower(), str2.lower()).ratio()
    
    def extract_with_training(self, text: str, field: str) -> str:
        """Extract field value using trained patterns and base models"""
        if not text.strip():
            return ""
        
        # Method 1: Use learned patterns
        if self.is_trained and field in self.field_patterns:
            pattern_result = self._extract_using_patterns(text, field)
            if pattern_result:
                return pattern_result
        
        # Method 2: Enhanced Q&A with context
        qa_result = self._extract_using_enhanced_qa(text, field)
        if qa_result:
            return qa_result
        
        # Method 3: Fallback extraction
        return self._extract_fallback(text, field)
    
    def _extract_using_patterns(self, text: str, field: str) -> str:
        """Extract using learned patterns from training data"""
        if field not in self.field_patterns:
            return ""
        
        patterns = self.field_patterns[field]
        if not patterns['contexts']:
            return ""
        
        # Find most similar context
        best_match = ""
        best_similarity = 0
        
        text_words = text.lower().split()
        
        for i, context in enumerate(patterns['contexts']):
            context_words = context.lower().split()
            
            # Calculate similarity based on common words
            common_words = set(text_words) & set(context_words)
            if len(context_words) > 0:
                similarity = len(common_words) / len(context_words)
                
                if similarity > best_similarity and similarity > 0.3:
                    best_similarity = similarity
                    best_match = patterns['values'][i]
        
        return best_match
    
    def _extract_using_enhanced_qa(self, text: str, field: str) -> str:
        """Enhanced Q&A extraction with better context handling"""
        if 'qa' not in self.base_models:
            return ""
        
        questions = self.extraction_strategies[field]["questions"]
        
        # Limit context and improve preprocessing
        context = text[:1500] if len(text) > 1500 else text
        context = ' '.join(context.split())  # Clean whitespace
        
        best_answer = ""
        best_score = 0
        
        for question in questions:
            try:
                result = self.base_models['qa'](
                    question=question,
                    context=context
                )
                
                if result['score'] > best_score and result['score'] > 0.15:
                    # Additional validation for the answer
                    answer = result['answer'].strip()
                    if self._is_valid_answer(answer, field):
                        best_answer = answer
                        best_score = result['score']
                        
            except Exception as e:
                continue
        
        return best_answer
    
    def _is_valid_answer(self, answer: str, field: str) -> bool:
        """Validate if extracted answer makes sense for the field"""
        if not answer or len(answer) < 2:
            return False
        
        if field == "Agreement Value":
            # Should contain numbers or currency symbols
            return any(char.isdigit() or char in ['$', ',', '.'] for char in answer)
        
        elif field in ["Agreement Start Date", "Agreement End Date"]:
            # Should look like a date
            try:
                parser.parse(answer, fuzzy=True)
                return True
            except:
                return any(char.isdigit() for char in answer)
        
        elif field == "Renewal Notice (Days)":
            # Should contain numbers
            return any(char.isdigit() for char in answer)
        
        elif field in ["Party One", "Party Two"]:
            # Should be a reasonable length for organization name
            return 3 <= len(answer) <= 100
        
        return True
    
    def _extract_fallback(self, text: str, field: str) -> str:
        """Fallback extraction method"""
        # This is a simplified fallback - in practice, you might use more sophisticated methods
        if field == "Party One":
            # Extract first organization-like entity
            entities = self._simple_entity_extraction(text)
            if entities:
                return entities[0]
        
        elif field == "Party Two":
            # Extract second organization-like entity
            entities = self._simple_entity_extraction(text)
            if len(entities) > 1:
                return entities[1]
        
        return ""
    
    def _simple_entity_extraction(self, text: str) -> List[str]:
        """Simple entity extraction as fallback"""
        if 'ner' not in self.base_models:
            return []
        
        try:
            entities = self.base_models['ner'](text[:1000])
            organizations = []
            
            for entity in entities:
                if (entity.get('entity_group') == 'ORG' or 
                    'ORG' in entity.get('entity', '') or
                    entity.get('label') == 'ORG'):
                    
                    org_name = entity['word'].strip()
                    if len(org_name) > 2:
                        organizations.append(org_name)
            
            return organizations[:3]  # Return top 3
        except:
            return []
    
    def extract_metadata(self, text: str) -> Dict[str, str]:
        """Extract all metadata fields"""
        metadata = {}
        
        for field in self.extraction_strategies.keys():
            if self.is_trained:
                extracted_value = self.extract_with_training(text, field)
            else:
                extracted_value = self._extract_using_enhanced_qa(text, field)
            
            metadata[field] = extracted_value
        
        return metadata

def fuzzy_match_score(str1: str, str2: str) -> float:
    """Calculate fuzzy matching score between two strings"""
    if not str1 or not str2:
        return 0.0
    
    # Normalize strings
    s1 = str1.lower().strip()
    s2 = str2.lower().strip()
    
    if s1 == s2:
        return 1.0
    
    # Use sequence matcher for fuzzy matching
    return SequenceMatcher(None, s1, s2).ratio()

def calculate_smart_recall(extracted_results: List[Dict], ground_truth_csv_path: str, threshold: float = 0.8):
    """Calculate recall with fuzzy matching for more realistic evaluation"""
    try:
        if not os.path.exists(ground_truth_csv_path):
            st.error(f"Ground truth CSV not found: {ground_truth_csv_path}")
            return None, None
        
        ground_truth_df = pd.read_csv(ground_truth_csv_path)
        
        fields_to_evaluate = [
            "Agreement Value", "Agreement Start Date", "Agreement End Date", 
            "Renewal Notice (Days)", "Party One", "Party Two"
        ]
        
        recall_metrics = {}
        detailed_results = []
        
        for field in fields_to_evaluate:
            true_positives = 0
            total_ground_truth = 0
            field_details = []
            
            for _, gt_row in ground_truth_df.iterrows():
                filename = gt_row.get('filename', '')
                gt_value = str(gt_row.get(field, '')).strip()
                
                if not gt_value or gt_value.lower() in ['nan', 'none', '']:
                    continue
                
                total_ground_truth += 1
                
                # Find corresponding extracted result
                extracted_value = ""
                for result in extracted_results:
                    if result['Filename'] == filename:
                        extracted_value = str(result.get(field, '')).strip()
                        break
                
                # Calculate fuzzy match score
                match_score = fuzzy_match_score(gt_value, extracted_value)
                is_match = match_score >= threshold
                
                if is_match:
                    true_positives += 1
                
                field_details.append({
                    'Filename': filename,
                    'Ground Truth': gt_value,
                    'Extracted': extracted_value,
                    'Match Score': f"{match_score:.2f}",
                    'Match': is_match
                })
            
            recall = true_positives / total_ground_truth if total_ground_truth > 0 else 0
            
            recall_metrics[field] = {
                'recall': recall,
                'true_positives': true_positives,
                'total_ground_truth': total_ground_truth,
                'details': field_details
            }
            
            detailed_results.extend([{
                'Field': field,
                'Filename': detail['Filename'],
                'Ground_Truth': detail['Ground Truth'],
                'Extracted_Value': detail['Extracted'],
                'Match_Score': detail['Match Score'],
                'Match': detail['Match']
            } for detail in field_details])
        
        return recall_metrics, detailed_results
        
    except Exception as e:
        st.error(f"Error calculating recall: {e}")
        return None, None

def display_recall_metrics(recall_metrics, title: str = "📊 Recall Metrics"):
    """Display recall metrics with improved visualization"""
    if not recall_metrics:
        return
    
    st.header(title)
    
    # Summary table
    summary_data = []
    for field, metrics in recall_metrics.items():
        summary_data.append({
            'Field': field,
            'Recall': f"{metrics['recall']:.2%}",
            'True Positives': metrics['true_positives'],
            'Total Ground Truth': metrics['total_ground_truth'],
            'Score': metrics['recall']  # For coloring
        })
    
    summary_df = pd.DataFrame(summary_data)
    
    st.subheader("📈 Per-Field Recall Summary")
    
    # Create colored dataframe
    def color_recall(val):
        score = float(val.rstrip('%')) / 100
        if score >= 0.8:
            return 'background-color: #d4edda'  # Green
        elif score >= 0.5:
            return 'background-color: #fff3cd'  # Yellow
        else:
            return 'background-color: #f8d7da'  # Red
    
    styled_df = summary_df[['Field', 'Recall', 'True Positives', 'Total Ground Truth']].style.map(
        color_recall, subset=['Recall']
    )
    
    st.dataframe(styled_df, use_container_width=True)
    
    # Overall metrics
    total_tp = sum(m['true_positives'] for m in recall_metrics.values())
    total_gt = sum(m['total_ground_truth'] for m in recall_metrics.values())
    overall_recall = total_tp / total_gt if total_gt > 0 else 0
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("🎯 Overall Recall", f"{overall_recall:.2%}")
    with col2:
        st.metric("✅ Total Matches", total_tp)
    with col3:
        st.metric("📝 Total Fields", total_gt)
    
    # Best and worst performing fields
    if recall_metrics:
        best_field = max(recall_metrics.items(), key=lambda x: x[1]['recall'])
        worst_field = min(recall_metrics.items(), key=lambda x: x[1]['recall'])
        
        col1, col2 = st.columns(2)
        with col1:
            st.success(f"🏆 Best: {best_field[0]} ({best_field[1]['recall']:.1%})")
        with col2:
            st.error(f"📉 Needs Work: {worst_field[0]} ({worst_field[1]['recall']:.1%})")

def process_files_with_training(directory_path: str, extractor: TrainableMetadataExtractor):
    """Process files using trained extractor"""
    results = []
    
    if not os.path.exists(directory_path):
        st.error(f"Directory not found: {directory_path}")
        return results
    
    files = [f for f in os.listdir(directory_path) if f.lower().endswith(('.docx', '.png', '.jpg', '.jpeg'))]
    
    if not files:
        st.warning(f"No supported files found in {directory_path}")
        return results
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for i, filename in enumerate(files):
        status_text.text(f"Processing: {filename}")
        filepath = os.path.join(directory_path, filename)
        
        try:
            # Extract text
            if filename.lower().endswith('.docx'):
                text = extractor.extract_text_from_docx(filepath)
            else:
                text = extractor.extract_text_from_image(filepath)
            
            if text.strip():
                # Extract metadata using trained model
                metadata = extractor.extract_metadata(text)
                metadata['Filename'] = filename
                metadata['File_Type'] = filename.split('.')[-1].upper()
                metadata['Text_Length'] = len(text)
                metadata['Model_Trained'] = extractor.is_trained
                results.append(metadata)
            else:
                st.warning(f"No text extracted from {filename}")
                
        except Exception as e:
            st.error(f"Error processing {filename}: {e}")
        
        progress_bar.progress((i + 1) / len(files))
    
    status_text.text("Processing complete!")
    return results

def main():
    st.sidebar.header("🎯 Training & Processing")
    
    # Load base models
    with st.spinner("Loading base AI models..."):
        base_models = load_base_models()
    
    if not base_models:
        st.error("Failed to load base models. Please check your internet connection.")
        return
    
    # Initialize extractor
    extractor = TrainableMetadataExtractor(base_models)
    
    # Step 1: Training Phase
    st.header("🏋️ Step 1: Training Phase")
    
    col1, col2 = st.columns(2)
    with col1:
        train_dir = st.text_input("📂 Training Directory Path:", value="train_dir", key="train_dir")
    with col2:
        train_csv = st.text_input("📄 Training CSV Path:", value="train.csv", key="train_csv")
    
    # Fuzzy matching threshold
    fuzzy_threshold = st.slider(
        "🎯 Fuzzy Matching Threshold (for recall calculation)", 
        min_value=0.5, max_value=1.0, value=0.8, step=0.05,
        help="Lower values allow more flexible matching. 0.8 means 80% similarity required."
    )
    
    if st.button("🚀 Train Model on Training Data", type="primary"):
        if train_dir and train_csv:
            with st.spinner("Training model on your data..."):
                extractor.train_on_data(train_dir, train_csv)
                
            if extractor.is_trained:
                st.session_state['trained_extractor'] = extractor
                st.session_state['fuzzy_threshold'] = fuzzy_threshold
        else:
            st.error("Please provide both training directory and CSV paths")
    
    # Show training status
    if hasattr(st.session_state, 'trained_extractor') and st.session_state['trained_extractor'].is_trained:
        st.success("✅ Model is trained and ready!")
        st.info(f"📊 Training data: {len(st.session_state['trained_extractor'].training_data)} documents")
    else:
        st.warning("⏳ Model not trained yet. Please train first using training data.")
    
    st.divider()
    
    # Step 2: Testing/Evaluation Phase
    st.header("🧪 Step 2: Testing & Evaluation")
    
    if hasattr(st.session_state, 'trained_extractor'):
        trained_extractor = st.session_state['trained_extractor']
        threshold = st.session_state.get('fuzzy_threshold', 0.8)
        
        col1, col2 = st.columns(2)
        with col1:
            test_dir = st.text_input("📂 Test Directory Path:", value="test_dir", key="test_dir")
        with col2:
            test_csv = st.text_input("📄 Test CSV Path:", value="test.csv", key="test_csv")
        
        if st.button("🔍 Process Test Data & Calculate Recall"):
            if test_dir and test_csv:
                # Process test files
                with st.spinner("Processing test files..."):
                    test_results = process_files_with_training(test_dir, trained_extractor)
                
                if test_results:
                    st.session_state['test_results'] = test_results
                    
                    # Calculate recall
                    with st.spinner("Calculating recall metrics..."):
                        recall_metrics, detailed_results = calculate_smart_recall(
                            test_results, test_csv, threshold
                        )
                    
                    if recall_metrics:
                        st.session_state['test_recall'] = recall_metrics
                        st.session_state['test_detailed'] = detailed_results
                        
                        # Display results
                        st.success(f"✅ Processed {len(test_results)} test files")
                        display_recall_metrics(recall_metrics, "🎯 Test Data Recall Metrics")
                        
                        # Show test results
                        st.subheader("📊 Test Extraction Results")
                        test_df = pd.DataFrame(test_results)
                        st.dataframe(test_df, use_container_width=True)
                        
                else:
                    st.error("No test files processed successfully")
            else:
                st.error("Please provide both test directory and CSV paths")
        
        # Training data evaluation (optional)
        st.subheader("📈 Training Data Evaluation (Optional)")
        if st.button("🔍 Evaluate on Training Data"):
            with st.spinner("Processing training files..."):
                train_results = process_files_with_training(train_dir, trained_extractor)
            
            if train_results:
                with st.spinner("Calculating training recall..."):
                    train_recall, train_detailed = calculate_smart_recall(
                        train_results, train_csv, threshold
                    )
                
                if train_recall:
                    st.session_state['train_recall'] = train_recall
                    display_recall_metrics(train_recall, "📚 Training Data Recall Metrics")
    
    else:
        st.info("👆 Please complete the training phase first")
    
    # Download Results
    if hasattr(st.session_state, 'test_results'):
        st.divider()
        st.header("💾 Download Results")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.session_state.get('test_results'):
                test_df = pd.DataFrame(st.session_state['test_results'])
                csv = test_df.to_csv(index=False)
                st.download_button(
                    label="📄 Download Test Results",
                    data=csv,
                    file_name=f"test_extraction_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv"
                )
        
        with col2:
            if st.session_state.get('test_detailed'):
                detailed_df = pd.DataFrame(st.session_state['test_detailed'])
                detailed_csv = detailed_df.to_csv(index=False)
                st.download_button(
                    label="🔍 Download Detailed Analysis",
                    data=detailed_csv,
                    file_name=f"detailed_recall_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv"
                )
        
        with col3:
            if st.session_state.get('test_recall'):
                recall_summary = []
                for field, metrics in st.session_state['test_recall'].items():
                    recall_summary.append({
                        'Field': field,
                        'Recall': metrics['recall'],
                        'True_Positives': metrics['true_positives'],
                        'Total_Ground_Truth': metrics['total_ground_truth']
                    })
                
                summary_df = pd.DataFrame(recall_summary)
                summary_csv = summary_df.to_csv(index=False)
                st.download_button(
                    label="📊 Download Recall Summary",
                    data=summary_csv,
                    file_name=f"recall_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv"
                )
    
    # Sidebar Information
    st.sidebar.header("🤖 Model Architecture")
    st.sidebar.info(
        """
        **Training Approach:**
        - Learns patterns from training data
        - Builds context vectors for each field
        - Adapts extraction strategies
        
        **Base Models:**
        - RoBERTa Question-Answering
        - BERT Named Entity Recognition
        - Sentence Transformers
        
        **Fuzzy Matching:**
        - Uses SequenceMatcher for similarity
        - Adjustable threshold for recall calculation
        - More realistic evaluation than exact matching
        """
    )
    
    st.sidebar.header("📊 Recall Calculation")
    st.sidebar.info(
        """
        **Smart Recall Formula:**
        
        Recall = True Positives / Total Ground Truth
        
        **Fuzzy Matching:**
        - Compares extracted vs ground truth values
        - Uses similarity scores (0.0 to 1.0)
        - Threshold determines what counts as "match"
        
        **Color Coding:**
        - Green: ≥80% recall
        - Yellow: 50-79% recall  
        - Red: <50% recall
        """
    )
    
    st.sidebar.header("Usage Instructions")
    st.sidebar.info(
        """
        **Step 1: Training**
        1. Provide training directory path
        2. Provide train.csv with metadata
        3. Click "Train Model" button
        4. Wait for training completion
        
        **Step 2: Testing**
        1. Provide test directory path
        2. Provide test.csv with ground truth
        3. Adjust fuzzy matching threshold
        4. Click "Process Test Data"
        5. View recall metrics and results
        
        **CSV Format Required:**
        ```
        filename,Agreement Value,Agreement Start Date,Agreement End Date,Renewal Notice (Days),Party One,Party Two
        contract1.docx,$50000,2023-01-01,2024-01-01,30,Company A,Company B
        ```
        
        **Supported Files:**
        - DOCX documents
        - PNG/JPG images (with OCR)
        """
    )

if __name__ == "__main__":
    main()